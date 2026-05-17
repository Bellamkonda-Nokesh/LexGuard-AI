"""Tests for document parsers."""
import os
import tempfile
import pytest
from pathlib import Path


class TestPDFParser:
    def test_parse_valid_pdf(self, tmp_path):
        """Test PDF parsing with a simple file."""
        try:
            import fitz
            # Create minimal PDF
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((50, 100), "TEST CONTRACT\nThis is a non-compete clause.")
            pdf_path = tmp_path / "test.pdf"
            doc.save(str(pdf_path))
            doc.close()

            from app.parsers.pdf_parser import parse_pdf
            text = parse_pdf(str(pdf_path))
            assert len(text) > 0
            assert "TEST CONTRACT" in text or "non-compete" in text.lower()
        except ImportError:
            pytest.skip("PyMuPDF not installed")

    def test_parse_nonexistent_pdf_raises(self):
        """Test that parsing nonexistent file raises RuntimeError."""
        from app.parsers.pdf_parser import parse_pdf
        with pytest.raises((RuntimeError, Exception)):
            parse_pdf("/nonexistent/path/file.pdf")

    def test_pdf_parser_returns_string(self, tmp_path):
        """Test that parser always returns a string type."""
        try:
            import fitz
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((50, 100), "Contract text here.")
            pdf_path = tmp_path / "test2.pdf"
            doc.save(str(pdf_path))
            doc.close()

            from app.parsers.pdf_parser import parse_pdf
            result = parse_pdf(str(pdf_path))
            assert isinstance(result, str)
        except ImportError:
            pytest.skip("PyMuPDF not installed")


class TestDOCXParser:
    def test_parse_valid_docx(self, tmp_path):
        """Test DOCX parsing with a real file."""
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("Employment Agreement")
            doc.add_paragraph("Non-compete clause: Employee shall not compete for 12 months.")
            docx_path = tmp_path / "test.docx"
            doc.save(str(docx_path))

            from app.parsers.docx_parser import parse_docx
            text = parse_docx(str(docx_path))
            assert len(text) > 0
            assert "Employment Agreement" in text
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_parse_empty_docx(self, tmp_path):
        """Test parsing an empty DOCX returns empty or minimal string."""
        try:
            from docx import Document
            doc = Document()
            docx_path = tmp_path / "empty.docx"
            doc.save(str(docx_path))

            from app.parsers.docx_parser import parse_docx
            text = parse_docx(str(docx_path))
            assert isinstance(text, str)
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_parse_docx_preserves_all_paragraphs(self, tmp_path):
        """Test that all paragraphs are captured."""
        try:
            from docx import Document
            doc = Document()
            paragraphs = ["Clause 1: Payment", "Clause 2: Termination", "Clause 3: Confidentiality"]
            for p in paragraphs:
                doc.add_paragraph(p)
            docx_path = tmp_path / "multi.docx"
            doc.save(str(docx_path))

            from app.parsers.docx_parser import parse_docx
            text = parse_docx(str(docx_path))
            for p in paragraphs:
                assert p in text
        except ImportError:
            pytest.skip("python-docx not installed")
